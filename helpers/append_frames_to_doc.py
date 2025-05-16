import os
import time
from docx import Document
from docx.shared import Inches

# Assuming paraphrase_transcript is defined elsewhere or imported
def paraphrase_transcript(text):
    # Dummy placeholder; replace with actual implementation or import
    return text.strip()

def append_frame_to_doc(frame_transcript_pairs, doc_path):
    doc = Document(doc_path) if os.path.exists(doc_path) else Document()
    for frame_path, transcript_text in frame_transcript_pairs:
        if os.path.exists(frame_path):
            doc.add_picture(frame_path, width=Inches(4))
            cleaned_text = paraphrase_transcript(transcript_text)
            # Uncomment below if you want to add timestamp info
            # timestamp = os.path.basename(frame_path).split('_')[-1].replace('.jpg', '')
            # doc.add_paragraph(f"Timestamp: {timestamp} sec")
            doc.add_paragraph(cleaned_text)
            time.sleep(5)  # Deliberate pause (e.g., for rate limiting)
    doc.save(doc_path)
    print(f"Document saved to: {doc_path}")


if __name__ == "__main__":
    # Example usage
    example_pairs = [
        ("path/to/frame_105_10.5.jpg", "This is a sample transcript text for the frame."),
        ("path/to/frame_210_21.0.jpg", "Another transcript segment related to the next frame.")
    ]
    output_doc = "output_document.docx"

    append_frame_to_doc(example_pairs, output_doc)
