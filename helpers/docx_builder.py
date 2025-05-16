from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Function to add a styled heading
def add_heading(doc, text, level, color=None):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    if color:
        run.font.color.rgb = RGBColor.from_string(color)

# Function to create a styled table
def create_table(doc, data, header_color="4F81BD"):
    table = doc.add_table(rows=len(data), cols=len(data[0]), style="Table Grid")
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), header_color)
                tc_pr = cell._element.get_or_add_tcPr()
                tc_pr.append(shading)

# Function to build the base Word document
def create_base_doc(project_name, process_summary, input_and_output):
    doc = Document()

    # Title Page
    doc.add_paragraph("\n\n", style=None)
    title = doc.add_paragraph()
    run = title.add_run(f"{project_name}\n")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.add_paragraph("\nProcess Definition Document (PDD)", style="Title")
    doc.add_page_break()

    # Table of Contents
    add_heading(doc, "TABLE OF CONTENTS", level=1, color="0D3B66")
    contents = [
        "1. Document Control",
        "   1.1 Version Control",
        "   1.2 Document Review and Sign-Off",
        "2. General Process Description",
        "   2.1 Document Purpose",
        "   2.2 Process Summary",
        "   2.3 Data Flow & Applications",
        "   2.4 Inputs and Outputs",
        "   2.5 High-Level Process Flow",
        "3. Handling Events – Known, Unknown, and Other",
        "   3.1 Known Exceptions",
        "   3.2 Unknown and Other Exceptions",
        "4. Step-by-Step Process Documentation"
    ]
    for item in contents:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Section 1: Document Control
    add_heading(doc, "1. DOCUMENT CONTROL", level=1, color="0D3B66")
    add_heading(doc, "1.1 VERSION CONTROL", level=2, color="0D3B66")
    create_table(doc, [["Version", "Date", "Description", "Author"], ["1.0", "", "", ""]])
    doc.add_paragraph("\n** Only document version checkpoint when shared with a primary stakeholder. **")
    add_heading(doc, "1.2 DOCUMENT REVIEW AND SIGN-OFF", level=2, color="0D3B66")
    create_table(doc, [["Name", "Business Role", "Action", "Date Reviewed"], ["", "", "", ""]])
    doc.add_page_break()

    # Section 2: General Process Description
    add_heading(doc, "2. GENERAL PROCESS DESCRIPTION", level=1, color="0D3B66")
    add_heading(doc, "2.1 DOCUMENT PURPOSE", level=2, color="0D3B66")
    doc.add_paragraph(
        "The purpose of the Process Design Document (PDD) is to capture the business-related details of "
        "the process being automated. It describes how the virtual workforce will operate. "
        "The PDD is a key input for the technical design of the solution."
    )
    doc.add_paragraph("Overall, the PDD ensures:")
    doc.add_paragraph("• Process requirements are captured in line with standards.")
    doc.add_paragraph("• It provides information on the process flow and detailed steps.")
    doc.add_paragraph("• Understanding of the expected results and objectives.")

    add_heading(doc, "2.2 PROCESS SUMMARY", level=2, color="0D3B66")
    doc.add_paragraph(process_summary)

    add_heading(doc, "2.3 DATA FLOW & APPLICATIONS", level=2, color="0D3B66")
    add_heading(doc, "2.3.1 APPLICATIONS INTERACTED WITH", level=3, color="000000")
    create_table(doc, [["Application", "Interface", "Key Operation / URL", "Comment"],
                       ["Testing", "1234", "ABCD", "XYZ"],
                       ["", "", "", ""],
                       ["", "", "", ""]])

    add_heading(doc, "2.4 INPUTS AND OUTPUTS", level=2, color="0D3B66")
    doc.add_paragraph(input_and_output)

    add_heading(doc, "2.5 HIGH LEVEL PROCESS FLOW", level=2, color="0D3B66")
    doc.add_paragraph("Flowchart will be added here dynamically.")
    doc.add_paragraph("\n\n")

    doc.add_page_break()

    # Section 3: Handling Events
    add_heading(doc, "2.6 BUSINESS RULES", level=2, color="0D3B66")
    doc.add_paragraph("N/A")

    add_heading(doc, "3. HANDLING EVENTS – KNOWN, UNKNOWN, AND OTHER", level=1, color="0D3B66")
    add_heading(doc, "3.1 KNOWN EXCEPTIONS", level=2, color="0D3B66")
    doc.add_paragraph("The following table contains process exceptions that this process may encounter:")
    create_table(doc, [["Event Code", "Description", "Action"],
                       ["", "", ""],
                       ["", "", ""],
                       ["", "", ""]])

    add_heading(doc, "3.2 UNKNOWN AND OTHER EXCEPTIONS", level=2, color="0D3B66")
    doc.add_paragraph(
        "‘Unknown Exceptions’ are exceptions that might occur during processing, that were expected and therefore not identified previously. "
        "These codes will be identified by ERROR_XYZ, and will be documented in the Technical Design Document (TDD)."
    )
    create_table(doc, [["Event Code", "Description", "Action"],
                       ["EVENT_001", "", ""],
                       ["EVENT_002", "", ""]])

    doc.add_page_break()

    # Section 4: Step-by-Step Process Documentation
    add_heading(doc, "4. STEP-BY-STEP PROCESS DOCUMENTATION", level=1, color="0D3B66")
    add_heading(doc, "4.1 STEP-BY-STEP DESCRIPTION", level=2, color="0D3B66")

    return doc
